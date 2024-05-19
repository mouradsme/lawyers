import json
from pydoc import doc
from django import forms
from django.shortcuts import render, redirect
from django.urls import reverse_lazy
from django.contrib.auth.views import LoginView, PasswordResetView, PasswordChangeView
from django.contrib import messages
from django.contrib.messages.views import SuccessMessageMixin
from django.views import View
from django.contrib.auth.decorators import login_required
from .forms import MessageForm
from .forms import RegisterForm, LawyerRegisterForm, LoginForm, UpdateUserForm, UpdateProfileForm,NotaryRegisterForm
from .models import NotaryClient, Profile, SubscriptionPlan
from .models import Documents, Notary, Profile, SubscriptionPlan, notaryclients, npaiements, nrdv, rdv
from .models import Lawyer ,User , Message,Notary, LawyerClient,lawyerclients
from .models import lawyerclients

from django.urls import reverse
from django.shortcuts import get_object_or_404
from itertools import groupby
from .models import User1, Message, LawyerClient
from django.http import HttpResponse
from datetime import date, datetime, timedelta
from django.utils import timezone
from datetime import timedelta
from django.db.models import Q
from .models import Aff,  Seance,  conseil, decision, paiement, tribunal,News
from django.http import JsonResponse 
from .forms import ETAPE_CHOICES, ETAT_CHOICES, JUDGMENT_CHOICES, NATURE_CHOICES, JugeForm,NewsForm
from django.db.models.signals import post_delete
from django.dispatch import receiver
from django.db.models import Sum
from django.contrib.auth.decorators import login_required
from django.contrib.admin.views.decorators import staff_member_required
from django.shortcuts import render
from django.core.files.base import ContentFile
from django.http import HttpResponseRedirect
def unauthorized_access(request):
    return render(request, 'unauthorized_access.html')
from functools import wraps

def lawyer_required(view_func):
    @wraps(view_func)
    def _wrapped_view(request, *args, **kwargs):
        # Check if the user is logged in and has the 'lawyer' role
        if request.user.is_authenticated and hasattr(request.user, 'role') and request.user.role == User.Role.LAWYER:
            return view_func(request, *args, **kwargs)
        else:
            # Render the error template with the error message
            error_message = "ليس لديك صلاحية الوصول إلى هذه الصفحة."
            return render(request, 'users/unauthorized_access.html', {'error_message': error_message})

    return _wrapped_view
def homesearch_results(request):
    if request.method == "GET":
        assistant_type = request.GET.get("assistant_type")
        practice_area = request.GET.get("practice_area")

        if assistant_type == "LAWYER":
            assistants = Lawyer.objects.filter(practice_areas__icontains=practice_area)
        elif assistant_type == "NOTARY":
            assistants = Notary.objects.filter(practice_areas__icontains=practice_area)
        else:
            assistants = []

        # Define the context dictionary
        context = {
            'assistants': assistants,
            # You can add more key-value pairs to the context as needed
        }

        return render(request, 'users/homesearch_results.html', context)
    else:
        return render(request, 'users/homesearch_results.html', {})

def registerall(request):
    # You can pass any necessary context data here

    return render(request, 'users/registerall.html')
def subscription_confirmation(request):
    # You can pass any necessary context data here

    return render(request, 'users/subscription_confirmation.html')
@lawyer_required
@login_required
def lawyerhome(request):
    
    # Get the current logged-in lawyer
    logged_in_lawyer = request.user

    # Get the current date and time
    current_date = timezone.now()

    # Print the current date in the console
    print("Current date:", current_date)

    # Get the start of the day
    start_of_day = timezone.datetime.combine(current_date.date(), timezone.datetime.min.time())

    # Get the end of the day
    end_of_day = timezone.datetime.combine(current_date.date(), timezone.datetime.max.time())

    # Calculate the count of seances for today for the logged-in lawyer
    seances_today = Seance.objects.filter(aff__lawyer=logged_in_lawyer, date__range=[start_of_day, end_of_day]).count()
    print("Seances for today:", seances_today)

    # Get the start of tomorrow
    start_of_tomorrow = current_date + timedelta(days=1)
    start_of_tomorrow = datetime.combine(start_of_tomorrow, datetime.min.time())

    # Calculate the count of seances for tomorrow for the logged-in lawyer
    seances_tomorrow = Seance.objects.filter(aff__lawyer=logged_in_lawyer, date__range=[start_of_tomorrow, start_of_tomorrow + timedelta(days=1)]).count()

    # Get the start of the week (assuming your week starts on Sunday)
    start_of_week = current_date - timedelta(days=current_date.weekday())
    start_of_week = datetime.combine(start_of_week, datetime.min.time())
 
    # Calculate the count of seances for the week for the logged-in lawyer
    seances_week = Seance.objects.filter(aff__lawyer=logged_in_lawyer, date__range=[start_of_week, start_of_week + timedelta(days=7)]).count()

    return render(request, 'users/lawyerhome.html', {
        'seances_today': seances_today,
        'seances_tomorrow': seances_tomorrow,
        'seances_week': seances_week,
    })
@login_required
def notaryhome(request):
    
    # Get the current logged-in lawyer
    logged_in_notary = request.user

    # Get the current date and time
    current_date = timezone.now()

    # Print the current date in the console
    print("Current date:", current_date)

    # Get the start of the day
    start_of_day = timezone.datetime.combine(current_date.date(), timezone.datetime.min.time())

    # Get the end of the day
    end_of_day = timezone.datetime.combine(current_date.date(), timezone.datetime.max.time())

    # Calculate the count of seances for today for the logged-in lawyer
    seances_today = Seance.objects.filter(aff__lawyer=logged_in_notary, date__range=[start_of_day, end_of_day]).count()
    print("Seances for today:", seances_today)

    # Get the start of tomorrow
    start_of_tomorrow = current_date + timedelta(days=1)
    start_of_tomorrow = datetime.combine(start_of_tomorrow, datetime.min.time())

    # Calculate the count of seances for tomorrow for the logged-in lawyer
    seances_tomorrow = Seance.objects.filter(aff__lawyer=logged_in_notary, date__range=[start_of_tomorrow, start_of_tomorrow + timedelta(days=1)]).count()

    # Get the start of the week (assuming your week starts on Sunday)
    start_of_week = current_date - timedelta(days=current_date.weekday())
    start_of_week = datetime.combine(start_of_week, datetime.min.time())
 
    # Calculate the count of seances for the week for the logged-in lawyer
    seances_week = Seance.objects.filter(aff__lawyer=logged_in_notary, date__range=[start_of_week, start_of_week + timedelta(days=7)]).count()

    return render(request, 'users/notaryhome.html', {
        'seances_today': seances_today,
        'seances_tomorrow': seances_tomorrow,
        'seances_week': seances_week,
    })

@login_required
def notaryhome(request):
    
    # Get the current logged-in lawyer
    logged_in_notary = request.user

    # Get the current date and time
    current_date = timezone.now()

    # Print the current date in the console
    print("Current date:", current_date)

    # Get the start of the day
    start_of_day = timezone.datetime.combine(current_date.date(), timezone.datetime.min.time())

    # Get the end of the day
    end_of_day = timezone.datetime.combine(current_date.date(), timezone.datetime.max.time())

    # Calculate the count of seances for today for the logged-in lawyer
    seances_today = Seance.objects.filter(aff__lawyer=logged_in_notary, date__range=[start_of_day, end_of_day]).count()
    print("Seances for today:", seances_today)

    # Get the start of tomorrow
    start_of_tomorrow = current_date + timedelta(days=1)
    start_of_tomorrow = datetime.combine(start_of_tomorrow, datetime.min.time())

    # Calculate the count of seances for tomorrow for the logged-in lawyer
    seances_tomorrow = Seance.objects.filter(aff__lawyer=logged_in_notary, date__range=[start_of_tomorrow, start_of_tomorrow + timedelta(days=1)]).count()

    # Get the start of the week (assuming your week starts on Sunday)
    start_of_week = current_date - timedelta(days=current_date.weekday())
    start_of_week = datetime.combine(start_of_week, datetime.min.time())
 
    # Calculate the count of seances for the week for the logged-in lawyer
    seances_week = Seance.objects.filter(aff__lawyer=logged_in_notary, date__range=[start_of_week, start_of_week + timedelta(days=7)]).count()

    return render(request, 'users/notaryhome.html', {
        'seances_today': seances_today,
        'seances_tomorrow': seances_tomorrow,
        'seances_week': seances_week,
    })






def Clients(request):
    clients = LawyerClient.objects.all()
    
    context = {'client': clients}
    return render(request, "users/Clients.html", context)
def affairepage(request):
   
    return render(request, "users/affaire.html")
@lawyer_required
@login_required  # Assure que l'utilisateur est connecté
def tablejugement(request):
    # Récupérer l'utilisateur connecté
    avocat = request.user

    # Filtrer les affaires liées à l'avocat connecté
    jugements = Aff.objects.filter(lawyer=avocat)

    context = {'Aff': jugements}
    return render(request, "users/tablejugement.html", context)


def detailed_view(request, Aff_id):
    affaire_instance = get_object_or_404(Aff, id=Aff_id)

    # Calculate the total payment for the case
    total_payment = paiement.objects.filter(aff=affaire_instance).aggregate(Sum('versement'))['versement__sum']

    # Calculate the remaining amount
    remaining_amount = affaire_instance.budjet - (total_payment if total_payment else 0)

    # Debugging print statements
    print(f"Total Payment: {total_payment}")
    print(f"Remaining Amount: {remaining_amount}")
    documents = Documents.objects.filter(aff=affaire_instance)

    seances = Seance.objects.filter(aff=affaire_instance)
    verdicts = decision.objects.filter(aff=affaire_instance)
    paiements = paiement.objects.filter(aff=affaire_instance)
    
    context = {
        'affaire': affaire_instance,
        'seances': seances,
        'verdicts': verdicts,
        'paiements': paiements,
        'documents': documents,

        'total_payment': total_payment if total_payment else 0,
        'remaining_amount': remaining_amount
    }
    return render(request, 'users/viewaffaire.html', context)

def printjuge(request, Aff_id):
    affaire_instance = get_object_or_404(Aff, id=Aff_id)
    context = {'affaire': affaire_instance, 'Aff_id': Aff_id}
    return render(request, 'users/printjuge.html', context)
@login_required
def printjt(request):
    # Retrieve the logged-in user's lawyer instance
    lawyer_instance = request.user.lawyer

    # Filter Aff instances based on the logged-in user's lawyer instance
    affaires = Aff.objects.filter(lawyer=lawyer_instance)

    return render(request, 'users/printjt.html', {'Aff': affaires})
def printpaiement(request, pay_id):
    # Retrieve the paiement instance with the given pay_id
    paiement_instance = get_object_or_404(paiement, id=pay_id)
    
    # Retrieve the aff instance associated with the paiement
    aff_instance = paiement_instance.aff
    
    # Pass both instances to the template context
    context = {'paiement': paiement_instance, 'affaire': aff_instance}
    
    # Render the template with the context
    return render(request, 'users/paiementpdf.html', context)
def printseance(request, sea_id):
    # Retrieve the paiement instance with the given pay_id
    seance_instance = get_object_or_404(Seance, id=sea_id)
    
    # Retrieve the aff instance associated with the paiement
    aff_instance = seance_instance.aff
    
    # Pass both instances to the template context
    context = {'seance': seance_instance, 'affaire': aff_instance}
    
    # Render the template with the context
    return render(request, 'users/PrintSeance.html', context)
def seances(request, Aff_id):
    affaire_instance = get_object_or_404(Aff, id=Aff_id)
    seances = Seance.objects.filter(aff=affaire_instance)
    context = {'seances': seances, 'Aff_id': Aff_id}  # Inclure Aff_id dans le contexte
    return render(request, 'users/seances.html', context)

def add_seance(request, Aff_id):
    affaire_instance = get_object_or_404(Aff, id=Aff_id)

    if request.method == "POST":
        st = request.POST.get('stitre')
        sd = request.POST.get('sdate')
        se = request.POST.get('sdetail')

        

        try:
            Seance.objects.create(aff=affaire_instance, titre=st, date=sd, remarques=se)
            messages.success(request, '  !  تمت اضافة الجلسة بنجاح ')
            # Redirect to the same page to prevent form resubmission on refresh
            return redirect('viewaffaire', Aff_id=affaire_instance.id)
        except Exception as e:
            messages.error(request, f'Error: {e}')

    # If it's a GET request or form submission fails, render the same page
    seances = Seance.objects.filter(aff=affaire_instance)
    context = {'affaire': affaire_instance, 'seances': seances}
    return render(request, "seances.html", context)
def updateseance(request, seance_id):
    seance = get_object_or_404(Seance, id=seance_id)
    if request.method == 'POST':
        titre = request.POST.get('stitre')
        date = request.POST.get('sdate')
        remarques = request.POST.get('sdetail')

        # Update only the fields that are modified
        if titre:
            seance.titre = titre
       
        if date:
            seance.date = date
        if remarques:
            seance.remarques = remarques
        seance.save()

        messages.success(request, 'تم تحديث الجلسة بنجاح')
        return redirect('seances', Aff_id=seance.aff_id)  # Corrected to use seance.aff_id
    
    return render(request, 'users/updateseance.html', {'seance': seance})

@lawyer_required
def add_verdict(request, Aff_id):
    affaire_instance = get_object_or_404(Aff, id=Aff_id)

    if request.method == "POST":
        va = request.POST.get('vdate')
        ve = request.POST.get('vdetail')

        try:
            decision.objects.create(aff=affaire_instance, date=va, detail=ve)
            messages.success(request, '!  تمت اضافة الحكم بنجاح ')
            return redirect('viewaffaire', Aff_id=affaire_instance.id)
        except Exception as e:
            messages.error(request, f'Error: {e}')

    # Debugging print statements
    print(f"Affaire ID: {Aff_id}")
    print(f"Verdict Date: {va}")
    print(f"Verdict Detail: {ve}")

    # Fetch the updated verdicts
    verdicts = decision.objects.filter(aff=affaire_instance)
    context = {'affaire': affaire_instance, 'verdicts': verdicts}
    return render(request, "users/verdict.html", context)


@lawyer_required
def add_paiement(request, Aff_id):
    affaire_instance = get_object_or_404(Aff, id=Aff_id)

    if request.method == "POST":
        pa = request.POST.get('pdate')
        ve = request.POST.get('versement')

        try:
            # Convert the date string to a datetime object
            pa_datetime = datetime.strptime(pa, '%Y-%m-%dT%H:%M')
            paiement.objects.create(aff=affaire_instance, versement=ve, date=pa_datetime)
            messages.success(request, '  !تمت اضافة الدفعة بنجاح')
            return redirect('viewaffaire', Aff_id=affaire_instance.id)
        except Exception as e:
            messages.error(request, f'Error: {e}')

    # Re-fetch paiements after adding a new payment
    paiements = paiement.objects.filter(aff=affaire_instance)
    context = {'affaire': affaire_instance, 'paiements': paiements}
    return render(request, "users/paiements.html", context)


def add_document(request, Aff_id):
    affaire_instance = get_object_or_404(Aff, id=Aff_id)

    if request.method == "POST":
        title = request.POST.get('title')
        file = request.FILES.get('file')

        try:
            Documents.objects.create(aff=affaire_instance, title=title, file=file)
            messages.success(request, 'تمت إضافة الملف بنجاح')
            return redirect('viewaffaire', Aff_id=affaire_instance.id)
        except Exception as e:
            messages.error(request, f'Error: {e}')

    # Re-fetch documents after adding a new document
    documents = Documents.objects.filter(aff=affaire_instance)
    context = {'affaire': affaire_instance, 'documents': documents}
    return render(request, "users/viewaffaire.html", context)

def document_list(request, Aff_id):
    affaire_instance = get_object_or_404(Aff, id=Aff_id)
    documents = Documents.objects.filter(aff=affaire_instance)
    context = {'Documents': documents, 'Aff_id': Aff_id}
    return render(request, 'viewaffaire.html', context)

@login_required
def display_clients(request):
    # Retrieve the authenticated user's lawyer instance
    lawyer = request.user.lawyer
    
    # Filter clients based on the associated lawyer
    clients = lawyer.lawyerclients_set.all()
    
    # Pass the client data to the template
    return render(request, 'users/Clients.html', {'clients': clients})


@login_required
def display_nclients(request):
    try:
        # Retrieve the authenticated user's notary instance
        notaire = request.user.notary
    except Notary.DoesNotExist:
        # Handle the case where the user doesn't have a notary
        notaire = None
    
    if notaire:
        # Filter clients based on the associated notary
        clients = notaire.notaryclients_set.all()
    else:
        clients = []
    
    # Pass the client data to the template
    return render(request, 'users/nclients.html', {'clients': clients})
    

@login_required
def pclients(request):
    if request.method == "POST":
        try:
            # Get form data
            fn = request.POST.get('Prenom')
            sn = request.POST.get('Nom')
            da = request.POST.get('date')
            em = request.POST.get('email')
            mo = request.POST.get('Numero')
            ad = request.POST.get('adresse')
            nc = request.POST.get('carte')
            dc = request.POST.get('cdate')

            # Get the authenticated user's lawyer instance
            lawyer = request.user.lawyer  # Assuming request.user is authenticated and associated with a Lawyer instance

            # Create a new lawyerclients instance and associate it with the lawyer
            lawyerclient = lawyerclients.objects.create(
                lawyer=lawyer,  # Associate with the lawyer
                Fname=fn,
                Sname=sn,
                date=da,
                email=em,
                mobile=mo,
                address=ad,
                numcarte=nc,
                datecarte=dc
            )
            messages.success(request, 'Client added successfully!')
            return redirect('Clients')
        except Exception as e:
            # If an error occurs during database operation, display an error message
            messages.error(request, f'Error: {e}')
    return render(request, "users/Clients.html")
   

@login_required
def pnclients(request):
    if request.method == "POST":
        try:
            # Get form data
            fn = request.POST.get('Prenom')
            sn = request.POST.get('Nom')
            da = request.POST.get('date')
            em = request.POST.get('email')
            mo = request.POST.get('Numero')
            ad = request.POST.get('adresse')
            nc = request.POST.get('carte')
            dc = request.POST.get('cdate')

            # Get the authenticated user's notary instance
            notaire = request.user.notary  # Assuming request.user is authenticated and associated with a Notary instance

            # Create a new notaryclients instance and associate it with the notary
            notaryclient = notaryclients.objects.create(
                notary=notaire,  # Associate with the notary
                Fname=fn,
                Sname=sn,
                date=da,
                email=em,
                mobile=mo,
                address=ad,
                numcarte=nc,
                datecarte=dc
            )
            messages.success(request, 'Client added successfully!')
            return redirect('nclients')
        except Exception as e:
            # If an error occurs during database operation, display an error message
            messages.error(request, f'Error: {e}')
    return render(request, "users/nclients.html")  

#def delete_client(request,cid):
  #  Client = LawyerClient.objects.get(id = cid)
  #  Client.delete()
def delete_nclient(request, cid):
    try:
        client = get_object_or_404(notaryclients, id=cid)
        client.delete()
        messages.success(request, 'تم حذف العميل بنجاح.')  # Add success message in Arabic
    except notaryclients.DoesNotExist:
        messages.error(request, 'العميل غير موجود.')  # Add error message in Arabic
    
    # Redirect back to the same page
    return HttpResponseRedirect(request.META.get('HTTP_REFERER'))
def delete_client(request, cid):
    try:
        client = get_object_or_404(lawyerclients, id=cid)
        client.delete()
        messages.success(request, 'تم حذف العميل بنجاح.')  # Add success message in Arabic
    except lawyerclients.DoesNotExist:
        messages.error(request, 'العميل غير موجود.')  # Add error message in Arabic
    
    # Redirect back to the same page
    return HttpResponseRedirect(request.META.get('HTTP_REFERER'))


def delete_doc(request, cid):
    try:
        doc = get_object_or_404(Documents, id=cid)
        doc.delete()
        messages.success(request, 'تم حذف المستند بنجاح.')  # Add success message in Arabic
    except Documents.DoesNotExist:
        # Handle the case where the document does not exist
        pass
    
    # Redirect back to the same page
    return HttpResponseRedirect(request.META.get('HTTP_REFERER'))


    return redirect('Clients')  
def update_client(request, client_id):
    client = get_object_or_404(lawyerclients, pk=client_id)
    if request.method == 'POST':
        # Update client information based on POST data
        client.Fname = request.POST.get('Fname')
        client.Sname = request.POST.get('Sname')
        client.date = request.POST.get('date')
        client.email = request.POST.get('email')
        client.mobile = request.POST.get('mobile')
        client.address = request.POST.get('address')
        client.numcarte = request.POST.get('numcarte')
        client.datecarte = request.POST.get('datecarte')
        client.save()
        messages.success(request, "تم تحديث معلومات العميل بنجاح.")

        return redirect('Clients')  # Redirect to client list page after update
    return render(request, 'users/upclient.html', {'client': client})
def update_nclient(request, client_id):
    client = get_object_or_404(notaryclients, pk=client_id)
    if request.method == 'POST':
        # Mettre à jour les informations du client en fonction des données POST
        client.Fname = request.POST.get('Fname')
        client.Sname = request.POST.get('Sname')
        client.date = request.POST.get('date')
        client.email = request.POST.get('email')
        client.mobile = request.POST.get('mobile')
        client.address = request.POST.get('address')
        client.numcarte = request.POST.get('numcarte')
        client.datecarte = request.POST.get('datecarte')
        client.save()
        # Rediriger vers la page de liste des clients après la mise à jour
        messages.success(request, "تم تحديث معلومات العميل بنجاح.")
        return redirect('nclients')  
    return render(request, 'users/upnclient.html', {'client': client})

@login_required
def rendezvous_list(request):
    current_date = timezone.now()  # Date et heure courantes
    lawyer = request.user.lawyer  # Get the authenticated user's lawyer instance
    upcoming_rendezvous = rdv.objects.filter(lawyer=lawyer, date__gte=current_date).order_by('date')
    past_rendezvous = rdv.objects.filter(lawyer=lawyer, date__lt=current_date).order_by('-date')
    return render(request, 'users/rdv.html', {'upcoming_rendezvous': upcoming_rendezvous, 'past_rendezvous': past_rendezvous})

@login_required
def nrendezvous_list(request):
    current_date = timezone.now()  # Date et heure courantes
    notary = request.user.notary  # Get the authenticated user's lawyer instance
    upcoming_rendezvous = nrdv.objects.filter(notary=notary, date__gte=current_date).order_by('date')
    past_rendezvous = nrdv.objects.filter(notary=notary, date__lt=current_date).order_by('-date')
    return render(request, 'users/nrdv.html', {'upcoming_rendezvous': upcoming_rendezvous, 'past_rendezvous': past_rendezvous})

@login_required
def prdv(request):
    if request.method == "POST":
        try:
            # Get form data
            fn = request.POST.get('nomprenom')
            da = request.POST.get('date')
            rm = request.POST.get('remarques')


            # Get the authenticated user's lawyer instance
            lawyer = request.user.lawyer  # Assuming request.user is authenticated and associated with a Lawyer instance

            # Create a new lawyerclients instance and associate it with the lawyer
            rendezvous = rdv.objects.create(
                lawyer=lawyer,  # Associate with the lawyer
                fullname=fn,
                date=da,
                remarques=rm,

            )
            messages.success(request, 'rendezvous added successfully!')
            return redirect('rdv')
        except Exception as e:
            # If an error occurs during database operation, display an error message
            messages.error(request, f'Error: {e}')
    return render(request, "users/rdv.html")
@login_required
def pnrdv(request):
    if request.method == "POST":
        try:
            # Get form data
            fn = request.POST.get('nomprenom')
            da = request.POST.get('date')
            rm = request.POST.get('remarques')


            # Get the authenticated user's lawyer instance
            notary = request.user.notary  # Assuming request.user is authenticated and associated with a Lawyer instance

            # Create a new lawyerclients instance and associate it with the lawyer
            rendezvous = nrdv.objects.create(
                notary=notary,  # Associate with the lawyer
                fullname=fn,
                date=da,
                remarques=rm,

            )
            messages.success(request, 'rendezvous added successfully!')
            return redirect('nrdv')
        except Exception as e:
            # If an error occurs during database operation, display an error message
            messages.error(request, f'Error: {e}')
    return render(request, "users/nrdv.html")


def delete_rdv(request,cid):
    try:
        rdvs = rdv.objects.get(id=cid)  # Rename the variable to avoid conflict
        rdvs.delete()
        messages.success(request, '! تمت الحذف بنجاح')
    except rdv.DoesNotExist:
        messages.error(request, 'العنصر المطلوب غير موجود')
    except Exception as e:
        messages.error(request, f'Error: {e}')
    
    return redirect('rdv')



def delete_nrdv(request,cid):
    try:
        rdvs = nrdv.objects.get(id=cid)  # Rename the variable to avoid conflict
        rdvs.delete()
        messages.success(request, '! تمت الحذف بنجاح')
    except nrdv.DoesNotExist:
        messages.error(request, 'العنصر المطلوب غير موجود')
    except Exception as e:
        messages.error(request, f'Error: {e}')
    
    return redirect('nrdv')


def update_rendezvous(request, rdv_id):
    rendezvous = get_object_or_404(rdv, id=rdv_id)
    
    if request.method == "POST":
        # Get the data from the request
        nomprenom = request.POST.get('edit_nomprenom')
        date = request.POST.get('edit_date')
        remarques = request.POST.get('edit_remarques')
        
        # Update the rendezvous object
        rendezvous.fullname = nomprenom
        rendezvous.date = date
        rendezvous.remarques = remarques
        rendezvous.save()
        return redirect('rdv')
    
    # Fetch the updated rendezvous again to ensure the latest data is displayed
    rendezvous = get_object_or_404(rdv, id=rdv_id)
    
    return render(request, 'users/uprdv.html', {'rdv': rendezvous})


def update_nrendezvous(request, rdv_id):
    rendezvous = get_object_or_404(nrdv, id=rdv_id)
    
    if request.method == "POST":
        # Get the data from the request
        nomprenom = request.POST.get('edit_nomprenom')
        date = request.POST.get('edit_date')
        remarques = request.POST.get('edit_remarques')
        
        # Update the rendezvous object
        rendezvous.fullname = nomprenom
        rendezvous.date = date
        rendezvous.remarques = remarques
        rendezvous.save()
        return redirect('nrdv')
    
    # Fetch the updated rendezvous again to ensure the latest data is displayed
    rendezvous = get_object_or_404(nrdv, id=rdv_id)
    
    return render(request, 'users/upnrdv.html', {'nrdv': rendezvous})

def delete_seance(request, cid):
    try:
        seance_instance = Seance.objects.get(id=cid)
        seance_instance.delete()
        messages.success(request, '! تمت الحذف بنجاح')
    except Seance.DoesNotExist:
        messages.error(request, 'العنصر المطلوب غير موجود')
    except Exception as e:
        messages.error(request, f'خطأ: {e}')
    
    # Redirect back to the same page
    return HttpResponseRedirect(request.META.get('HTTP_REFERER', '/'))


def delete_pay(request, cid):
    try:
        pay_instance = paiement.objects.get(id=cid)  # Rename the variable to avoid conflict
        pay_instance.delete()
        messages.success(request, '! تمت الحذف  بنجاح')
    except paiement.DoesNotExist:
        messages.error(request, 'العنصر المطلوب غير موجود')
    except Exception as e:
        messages.error(request, f'Error: {e}')
    
     # Redirect back to the same page
    return HttpResponseRedirect(request.META.get('HTTP_REFERER', '/'))


def tribunal_list(request):
    tribunals = tribunal.objects.all()
    return render(request, 'users/tribunals.html', {'tribunals': tribunals})

@lawyer_required
@login_required  # Ensure the user is logged in
def jugements(request):
    if request.method == "POST":
        try:
            # Get form data
            type = request.POST.get('type')
            numeroaffaire = request.POST.get('numeroaffaire')
            etape = request.POST.get('etape')
            etat = request.POST.get('etat')
            conseil_id = request.POST.get('conseil')
            tribunal_id = request.POST.get('tribunal')
            sujeta = request.POST.get('sujeta')
            datea = request.POST.get('datea')
            budjeta = request.POST.get('budjeta')
            adver = request.POST.get('adver')
            adre = request.POST.get('adre')
            avocata = request.POST.get('avocata')
            tclient = request.POST.get('tclient')
            nomclient = request.POST.get('nomclient') 
            num = request.POST.get('num')
            details = request.POST.get('details') # Get client's name from POST data

            # Get the authenticated user's lawyer instance
            lawyer = request.user.lawyer

            # Fetch the conseil and tribunal instances based on the provided IDs
            conseil_instance = conseil.objects.get(pk=conseil_id)
            tribunal_instance = tribunal.objects.get(pk=tribunal_id)

            # Create a new Aff instance and associate it with the lawyer
            aff_instance = Aff.objects.create(
                lawyer=lawyer,
                taffaire=type,
                Naff=numeroaffaire,
                etapeaffaire=etape,
                etataffaire=etat,
                conseil_associated=conseil_instance,
                tribunal_associated=tribunal_instance,
                sujet=sujeta,
                date=datea,
                budjet=budjeta,
                adversaire=adver,
                adresse=adre,
                avocat=avocata,
                typeclient=tclient,
                Nomclient=nomclient ,
                details=details,
                tel=num,

            )
            messages.success(request, 'Judgment added successfully!')
            return redirect('jugements')
        except Exception as e:
            # If an error occurs during database operation, display an error message
            messages.error(request, f'Error: {e}')
            return HttpResponseRedirect(request.path_info)  # Redirect to the same page
    else:
        # Retrieve LawyerClient instances associated with the logged-in lawyer
        lawyer = request.user.lawyer
        clients = lawyerclients.objects.filter(lawyer=lawyer)
        
        # Initialize nomclient with the first client's name if available
        initial_nomclient = f"{clients.first().Fname} {clients.first().Sname}" if clients.exists() else None
        # Set initial value for nomclient and update choices
        form = JugeForm(initial={'nomclient': initial_nomclient})
        form.fields['nomclient'].choices = [(f"{client.Fname} {client.Sname}", f"{client.Fname} {client.Sname}") for client in clients]

        return render(request, "users/jugements.html", {'form': form, 'clients': clients})
def upjugements(request, jugement_id):
    jugement = get_object_or_404(Aff, id=jugement_id)
    
    if request.method == "POST":
        form = JugeForm(request.POST)
        if form.is_valid():
            # Get the data from the form
            taffaire = form.cleaned_data['type']
            numeroaffaire = form.cleaned_data['numeroaffaire']
            etapeaffaire = form.cleaned_data['etape']
            etataffaire = form.cleaned_data['etat']
            conseil_associated = form.cleaned_data['conseil']
            tribunal_associated = form.cleaned_data['tribunal']
            sujet = form.cleaned_data['sujeta']
            date = form.cleaned_data['datea']
            budjet = form.cleaned_data['budjeta']
            adversaire = form.cleaned_data['adver']
            adresse = form.cleaned_data['adre']
            avocat = form.cleaned_data['avocata']
            typeclient = form.cleaned_data['tclient']
            nomclient = form.cleaned_data['nomclient']
            details = form.cleaned_data['details']
            tel = form.cleaned_data['num']
            
            # Update the judgment object with the form data
            jugement.taffaire = taffaire
            jugement.Naff = numeroaffaire
            jugement.etapeaffaire = etapeaffaire
            jugement.etataffaire = etataffaire
            jugement.conseil_associated = conseil_associated
            jugement.tribunal_associated = tribunal_associated
            jugement.sujet = sujet
            jugement.date = date
            jugement.budjet = budjet
            jugement.adversaire = adversaire
            jugement.adresse = adresse
            jugement.avocat = avocat
            jugement.typeclient = typeclient
            jugement.Nomclient = nomclient
            jugement.details = details
            jugement.tel = tel
            jugement.save()
            
            # Add a success message in Arabic
            messages.success(request, "تم تحديث البيانات بنجاح.")
            
            return redirect('tablejugement')
    else:
        form = JugeForm(initial={
            'type': jugement.taffaire,
            'numeroaffaire': jugement.Naff,
            'etape': jugement.etapeaffaire,
            'etat': jugement.etataffaire,
            'conseil': jugement.conseil_associated,
            'tribunal': jugement.tribunal_associated,
            'sujeta': jugement.sujet,
            'datea': jugement.date,
            'budjeta': jugement.budjet,
            'adver': jugement.adversaire,
            'adre': jugement.adresse,
            'avocata': jugement.avocat,
            'tclient': jugement.typeclient,
            'nomclient': jugement.Nomclient,
            'details': jugement.details,
            'num': jugement.tel,
        })
    
    return render(request, 'users/upjugement.html', {'form': form, 'judgment': jugement})
def delete_affaire(request, cid):
    try:
        affaire_instance = Aff.objects.get(id=cid)
        affaire_instance.delete()
        messages.success(request, '! تمت الحذف بنجاح')
    except Seance.DoesNotExist:
        messages.error(request, 'العنصر المطلوب غير موجود')
    except Exception as e:
        messages.error(request, f'خطأ: {e}')
    
    # Redirect back to the same page
    return HttpResponseRedirect(request.META.get('HTTP_REFERER', '/'))


@login_required
def tableseances(request):
    # Retrieve the logged-in user
    user = request.user
    
    # Check if the user is a lawyer
    if hasattr(user, 'lawyer'):
        # Retrieve the sessions associated with the logged-in lawyer
        seance_objects = Seance.objects.filter(aff__lawyer=user.lawyer)
        
        # Pass the filtered session objects to the template context
        return render(request, 'users/tableseances.html', {'seance_objects': seance_objects})
    else:
        # If the user is not a lawyer, redirect or display an error message
        # Redirect to an appropriate page or display an error message
        return HttpResponse("أنت لست محامياً.")

@lawyer_required
@login_required
def todayseances(request):
    # Get today's date
    today_date = date.today()

    # Get the logged-in lawyer
    lawyer = request.user.lawyer

    # Query Seance objects for the current day for the logged-in lawyer
    tseance_objects = Seance.objects.filter(aff__lawyer=lawyer, date__date=today_date)

    # Pass the filtered Seance objects to the template context
    return render(request, 'users/todayseances.html', {'tseance_objects': tseance_objects})
@lawyer_required
@login_required
def tomorrowseances(request):
    # Get tomorrow's date
    tomorrow_date = date.today() + timedelta(days=1)

    # Get the logged-in lawyer
    lawyer = request.user.lawyer

    # Query Seance objects for the next day for the logged-in lawyer
    tmseance_objects = Seance.objects.filter(aff__lawyer=lawyer, date__date=tomorrow_date)

    # Pass the filtered Seance objects to the template context
    return render(request, 'users/tomorowseances.html', {'tmseance_objects': tmseance_objects})
@lawyer_required
@login_required
def weekseances(request):
    # Get today's date
    today = date.today()
    # Calculate the start and end dates of the week
    start_of_week = today - timedelta(days=today.weekday())
    end_of_week = start_of_week + timedelta(days=6)
    
    # Get the logged-in lawyer
    lawyer = request.user.lawyer

    # Query Seance objects for the week for the logged-in lawyer
    week_seance_objects = Seance.objects.filter(aff__lawyer=lawyer, date__date__range=[start_of_week, end_of_week])

    # Pass the filtered Seance objects to the template context
    return render(request, 'users/weekseances.html', {'week_seance_objects': week_seance_objects})
@lawyer_required
def get_conseil_data(request):
    conseil_data = list(conseil.objects.values('id', 'name'))
    return JsonResponse(conseil_data, safe=False)
@lawyer_required
def get_tribunals(request):
    conseil_id = request.GET.get('conseil')
    tribunals = tribunal.objects.filter(college_id=conseil_id).values('id', 'name')
    options = [{'value': tribunal['id'], 'label': tribunal['name']} for tribunal in tribunals]
    return JsonResponse(options, safe=False)


#end lawyer----------------------------------------------------





def home(request):
    if request.user.is_authenticated:
        if hasattr(request.user, 'role'):
            if request.user.role == User.Role.LAWYER:
                return redirect('lawyerhome')
            elif request.user.role == User.Role.USER1:
                return redirect('home2')
            elif request.user.role == User.Role.NOTARY:  
                return redirect('notaryhome')
        else:
            news_articles = News.objects.all()[:5]  
            return render(request, 'users/home.html', {'news_articles': news_articles})
    else:
        return render(request, 'users/home.html')
    
def register_notary(request):
    if request.method == 'POST':
        form = NotaryRegisterForm(request.POST)
        if form.is_valid():
            form.save()  # Save the form data to create a new Notary object
            return redirect('Notaryall')  # Redirect to a success page or any other page
    else:
        form = NotaryRegisterForm()

    return render(request, 'users/register_notary.html', {'form': form})

def search_notaries(request):
    query = request.GET.get('query', '')  # Get the search query from the request

    # Filter notaries by username, practice_areas, and experience_years containing the query
    notaries = Notary.objects.filter(
        Q(username__icontains=query) |  # Search by username
        Q(practice_areas__icontains=query) |  # Search by practice_areas
        Q(experience_years__icontains=query)  # Search by experience_years
    ).distinct()

    # Pass the search results and query back to the template
    return render(request, 'users/Notaryall.html', {'notaries': notaries, 'query': query})


    
def search_lawyers(request):
    query = request.GET.get('query', '')  # Get the search query from the request

    # Filter lawyers by username, practice_areas, and experience_years containing the query
    lawyers = Lawyer.objects.filter(
        Q(username__icontains=query) |  # Search by username
        Q(practice_areas__icontains=query) |  # Search by practice_areas
        Q(experience_years__icontains=query)  # Search by experience_years
    ).distinct()

    # Pass the search results and query back to the lawall template
    return render(request, 'users/Lawall.html', {'lawyers': lawyers, 'query': query})
def Notaryall(request):
    return render(request, 'users/Notaryall.html')

def test(request):
    return render(request, 'users/test.html')   

def basetest(request):
    return render(request, 'users/basetest.html')

def home2(request):
    news_articles = News.objects.all()[:10]  # Display 10 latest news articles for home2
    return render(request, 'users/home.html', {'news_articles': news_articles})


def Lawall(request):
    return render(request, 'users/Lawall.html')
def chat(request):
    return render(request, 'users/index.html')
def contrats(request):
    return render(request, 'users/contrats.html') 

class RegisterView(View):
    form_class = RegisterForm
    initial = {'key': 'value'}
    template_name = 'users/register.html'

    def dispatch(self, request, *args, **kwargs):
        # will redirect to the home page if a user tries to access the register page while logged in
        if request.user.is_authenticated:
            return redirect(to='/')

        # else process dispatch as it otherwise normally would
        return super(RegisterView, self).dispatch(request, *args, **kwargs)

    def get(self, request, *args, **kwargs):
        form = self.form_class(initial=self.initial)
        return render(request, self.template_name, {'form': form})

    def post(self, request, *args, **kwargs):
        form = self.form_class(request.POST)
        
        if form.is_valid():
            user = form.save(commit=False)
            user.role = User.Role.USER1  # Automatically set the role to USER1
            user.save()
            #user = form.save(commit=False)
           # user.role = User.Role.USER1  # Assigning the role of USER1
            #user.set_password(form.cleaned_data['password1'])
           #soure  user.save()


            username = form.cleaned_data.get('username')
            messages.success(request, f'تم إنشاء حساب لـ {username} بنجاح')


            return redirect(to='login')

        return render(request, self.template_name, {'form': form})



class LawyerRegisterView(View):
    form_class = LawyerRegisterForm
    template_name = 'users/lawregister.html'

    def dispatch(self, request, *args, **kwargs):
    # will redirect to the home page if a user tries to access the register page while logged in
     if request.user.is_authenticated:
        return redirect(to='/')

    # else process dispatch as it otherwise normally would
     return super(self.__class__, self).dispatch(request, *args, **kwargs)

    def get(self, request, *args, **kwargs):
        form = self.form_class()
        return render(request, self.template_name, {'form': form})

    def post(self, request, *args, **kwargs):
        form = self.form_class(request.POST)

        if form.is_valid():
            
            user = form.save(commit=False)
            user.role = User.Role.LAWYER  # Automatically set the role to LAWYER
            password = form.cleaned_data['password1']  # Adjusted to password1
            user.set_password(password)  # Manually set the password
            
            user.save()
            
            Profile.objects.create(user=user)
            return redirect('login')

        return render(request, self.template_name, {'form': form})
    

class NotaryRegisterView(View):
    form_class = NotaryRegisterForm
    template_name = 'users/notary_register.html'

    def get(self, request, *args, **kwargs):
        form = self.form_class()
        return render(request, self.template_name, {'form': form})

    def post(self, request, *args, **kwargs):
        form = self.form_class(request.POST)

        if form.is_valid():
            user = form.save(commit=False)
            user.role = User.Role.NOTARY  # Set the role to NOTARY
            password = form.cleaned_data['password1']
            user.set_password(password)
            user.save()

            Profile.objects.create(user=user)
            return redirect('login')  # Redirect to login page upon successful registration

        return render(request, self.template_name, {'form': form})



def list_lawyers_with_profiles(request):
    lawyers_with_profiles = Lawyer.objects.select_related('profile').all()
    return render(request, 'users/lawyers_with_profiles.html', {'lawyers_with_profiles': lawyers_with_profiles})

def lawyer_register(request):
    if request.method == 'POST':
        form = LawyerRegisterForm(request.POST)
        if form.is_valid():
            form.save()
            username = form.cleaned_data.get('username')
            messages.success(request, f'تم إنشاء حساب لـ {username} بنجاح')

            return redirect('Lawall')
    else:
        form = LawyerRegisterForm()
    return render(request, 'users/lawyer_register.html', {'form': form})
# Class based view that extends from the built in login view to add a remember me functionality
class CustomLoginView(LoginView):
    form_class = LoginForm

    def form_valid(self, form):
        remember_me = form.cleaned_data.get('remember_me')

        if not remember_me:
            # set session expiry to 0 seconds. So it will automatically close the session after the browser is closed.
            self.request.session.set_expiry(0)

            # Set session as modified to force data updates/cookie to be saved.
            self.request.session.modified = True

        # else browser session will be as long as the session cookie time "SESSION_COOKIE_AGE" defined in settings.py
        return super(CustomLoginView, self).form_valid(form)


class ResetPasswordView(SuccessMessageMixin, PasswordResetView):
    template_name = 'users/password_reset.html'
    email_template_name = 'users/password_reset_email.html'
    subject_template_name = 'users/password_reset_subject'
    success_message = "We've emailed you instructions for setting your password, " \
                      "if an account exists with the email you entered. You should receive them shortly." \
                      " If you don't receive an email, " \
                      "please make sure you've entered the address you registered with, and check your spam folder."
    success_url = reverse_lazy('users-home')


class ChangePasswordView(SuccessMessageMixin, PasswordChangeView):
    template_name = 'users/change_password.html'
    success_message = "Successfully Changed Your Password"
    success_url = reverse_lazy('users-home')


@login_required
def profile(request):
    if request.method == 'POST':
        user_form = UpdateUserForm(request.POST, instance=request.user)
        profile_form = UpdateProfileForm(request.POST, request.FILES, instance=request.user.profile)

        if user_form.is_valid() and profile_form.is_valid():
            user_form.save()
            profile_form.save()
            messages.success(request, 'تم تحديث ملفك الشخصي بنجاح.')
            return redirect(to='users-profile')
    else:
        user_form = UpdateUserForm(instance=request.user)
        profile_form = UpdateProfileForm(instance=request.user.profile)

    return render(request, 'users/profile.html', {'user_form': user_form, 'profile_form': profile_form})



def send_message(request):
    if request.method == 'POST':
        recipient_id = request.POST.get('recipient_id')  # Assuming the ID of the recipient is passed in the request
        content = request.POST.get('content')  
        sender = request.user  # Assuming the current user is the sender
        
        # Create and save the message
        message = Message(sender=sender, recipient_id=recipient_id, content=content)
        message.save()
        
        messages.success(request, f'تم إرسال الرسالة إلى المستلم ذو الهوية {recipient_id} بنجاح')
        return redirect('all_lawyers')
    else:
        return redirect('all_lawyers')  # Redirect if not a POST request or if no recipient provided

@login_required
def send_message_ajax(request):
    if request.method == 'POST' and request.is_ajax():
        recipient_id = request.POST.get('recipient_id')
        content = request.POST.get('content')
        # Create message
        # You'll need to adjust this part according to your Message model
        message = create_message(sender_id=request.user.id, recipient_id=recipient_id, content=content)
        return JsonResponse({'status': 'success', 'message_id': message.id})
    return JsonResponse({'status': 'error'})

@login_required
def fetch_messages_ajax(request, sender_id):
    if request.is_ajax():
        # Fetch messages between current user and sender_id
        messages = Message.objects.filter(
            (Q(sender=request.user) & Q(recipient_id=sender_id)) |
            (Q(sender_id=sender_id) & Q(recipient=request.user))
        ).order_by('timestamp')
        # Serialize messages to JSON
        messages_data = [{'content': message.content, 'timestamp': message.timestamp} for message in messages]
        return JsonResponse({'status': 'success', 'messages': messages_data})
    return JsonResponse({'status': 'error'})

def create_message(sender_id, recipient_id, content):
    # Fetch the sender user instance
    try:
        sender = User.objects.get(id=sender_id)
    except User.DoesNotExist:
        raise ValueError("Sender with the provided ID does not exist")

    # Create a new message object
    message = Message.objects.create(
        sender_id=sender,
        recipient=recipient_id,
        content=content
    )
    # Optionally, you can perform additional operations here if needed
    # For example, sending notifications, updating counters, etc.




@login_required
def message_form(request, recipient_id):
    recipient = User.objects.get(pk=recipient_id)
    if request.method == 'POST':
        form = MessageForm(request.POST)
        if form.is_valid():
            content = form.cleaned_data['content']
            # Create a new Message instance and set its attributes
            message = Message.objects.create(
                sender_id=request.user,  # Set the sender_id to the current user instance
                recipient=recipient,
                content=content
            )
            messages.success(request, 'تم إرسال الرسالة بنجاح!')

            return redirect('inbox')
    else:
        form = MessageForm(initial={'recipient': recipient})
    return render(request, 'users/message_form.html', {'form': form, 'recipient': recipient})



#this is it 
def message_view(request):
    messages = Message.objects.all()  # Retrieve all messages
    return render(request, 'users/inboxnew.html', {'messages': messages})
def Lawall(request):
    
    lawyers = Lawyer.objects.all()
    return render(request, 'users/Lawall.html', {'lawyers': lawyers})


@login_required
def inboxnew(request):
    user = request.user
    
    # Initialize sender and recipient profile pictures
    sender_profile_picture = None
    recipient_profile_picture = None
    
    # Fetch messages sent by the current user
    sent_messages = Message.objects.filter(sender_id=user.id)
    grouped_messages = {}
    for message in sent_messages:
        grouped_messages.setdefault(message.recipient_id, []).append(message)

    # If the request is POST, process the form to send a new message
    if request.method == 'POST':
        form = MessageForm(request.POST)
        if form.is_valid():
            recipient_id = form.cleaned_data['recipient_id']
            content = form.cleaned_data['content']
            create_message(sender_id=user.id, recipient_id=recipient_id, content=content)
            messages.success(request, 'تم إرسال الرسالة بنجاح!')

            return redirect('inboxnew')
    else:
        form = MessageForm()

    if user.role == User.Role.USER1:
        # Fetch messages where the recipient is the current user and the sender's role is a lawyer
        received_messages = Message.objects.filter(recipient=user, sender_id__role=User.Role.LAWYER)
    elif user.role == User.Role.LAWYER:
        # Fetch users accepted by the lawyer
        accepted_clients = LawyerClient.objects.filter(lawyer=user.lawyer).values_list('user1', flat=True)
        # Filter messages sent by accepted clients to the lawyer
        received_messages = Message.objects.filter(sender_id__in=accepted_clients, recipient=user)
        
    # Add code to retrieve sender and recipient profile pictures
    if hasattr(request.user, 'profile'):
        sender_profile_picture = request.user.profile.avatar.url
        recipient_id = request.POST.get('recipient_id')
        if recipient_id:
            recipient = User.objects.get(pk=recipient_id)
            if hasattr(recipient, 'profile'):
                recipient_profile_picture = recipient.profile.avatar.url

    return render(request, 'users/inboxnew.html', {
        'grouped_messages': grouped_messages,
        'received_messages': received_messages,
        'form': form,
        'sender_profile_picture': sender_profile_picture,
        'recipient_profile_picture': recipient_profile_picture
    })




def lawyer_profile(request):
    # Fetch messages where the recipient is also the lawyer
    lawyer_messages = Message.objects.filter(recipient_id=request.user.id)

    # Check if the user is a lawyer
    if request.user.role == User.Role.LAWYER:
        # Get the lawyer's profile
        lawyer_profile = Lawyer.objects.get(id=request.user.id)
        return render(request, 'lawyer_profile.html', {'lawyer_profile': lawyer_profile})
    else:
        return render(request, 'lawyer_profile.html', {'lawyer_messages': lawyer_messages})

def message_detail(request, message_id):
    message = get_object_or_404(Message, id=message_id)
    recipient_lawyer = message.recipient  # This will fetch the recipient lawyermanalawyer
    return render(request, 'users/inboxnew.html', {'message': message, 'recipient_lawyer': recipient_lawyer})







# views.py



@login_required
def manage_clients(request):
    if request.user.role != User.Role.LAWYER:
        messages.error(request, "ليس لديك الصلاحية للوصول إلى هذه الصفحة.")

        return redirect('home2')

    lawyer = request.user.lawyer
    if lawyer is None:
        messages.error(request, "You are not registered as a lawyer.")
        return redirect('home2')

    if request.method == 'POST':
        selected_clients = request.POST.getlist('selected_clients')
        subscription = request.user.profile.subscription_plan  # Access subscription plan through the user's profile

        if subscription is None:
            messages.error(request, "ليس لديك خطة اشتراك نشطة.")

            return redirect('home2')

        max_accepted_clients = subscription.max_accepted_clients
        accepted_count = lawyer.accepted_clients.count()

        non_existent_clients = []

        for client_id in selected_clients:
            try:
                client = User1.objects.get(pk=client_id)
            except User1.DoesNotExist:
                non_existent_clients.append(client_id)
                continue

            if LawyerClient.objects.filter(lawyer=lawyer, user1=client).exists():
                messages.warning(request, f"العميل '{client.username}' مقبول بالفعل.")

            elif max_accepted_clients is not None and accepted_count >= max_accepted_clients:
                messages.error(request, "لقد وصلت إلى الحد الأقصى لعدد العملاء المقبولين.")

            else:
                LawyerClient.objects.create(lawyer=lawyer, user1=client)
                messages.success(request, f"تم قبول العميل '{client.username}' بنجاح.")

                # Send a message to the client
                content = "تم قبولك من قبل المحامي."
                Message.objects.create(sender_id=request.user, recipient=client, content=content)
                accepted_count += 1

        if non_existent_clients:
            messages.error(request, f"The following clients do not exist: {', '.join(non_existent_clients)}")

        return redirect('manage_clients')

    else:
        message_senders = Message.objects.filter(recipient=request.user)
        potential_clients = User1.objects.filter(id__in=message_senders.values('sender_id'))
        context = {'potential_clients': potential_clients}
        return render(request, 'users/manage_clients.html', context)



    


from django.utils import timezone
from datetime import timedelta
from django.shortcuts import render, redirect
from django.contrib import messages
from django.core.mail import send_mail
from django.template.loader import render_to_string
from django.utils.html import strip_tags
from .models import SubscriptionPlan
from django.contrib import messages
from django.shortcuts import render, redirect
from django.utils import timezone
from django.core.mail import send_mail
from django.template.loader import render_to_string
from django.utils.html import strip_tags
from datetime import timedelta
from .models import SubscriptionPlan

def subscribe(request):
    if request.method == 'POST':
        subscription_plan_name = request.POST.get('subscription_plan')
        if subscription_plan_name:
            try:
                new_subscription_plan = SubscriptionPlan.objects.get(name=subscription_plan_name)
            except SubscriptionPlan.DoesNotExist:
                messages.error(request, f"خطة الاشتراك '{subscription_plan_name}' غير موجودة.")
                return redirect('subscribe')
            
            # Check if the user has an active subscription
            if request.user.is_authenticated:
                user_profile = request.user.profile
                current_subscription_end_date = user_profile.subscription_end_date
                if current_subscription_end_date and current_subscription_end_date > timezone.now():
                    # Check if the new plan has a longer duration than the current one
                    if new_subscription_plan.duration > user_profile.subscription_plan.duration:
                        messages.error(request, "لديك بالفعل اشتراك نشط. لا يمكنك شراء اشتراك جديد حتى ينتهي الحالي.")
                        return redirect('subscribe')

                # Check if the user is trying to downgrade the plan
                if user_profile.subscription_plan and new_subscription_plan.price < user_profile.subscription_plan.price:
                    messages.error(request, "لا يمكنك تخفيض خطة الاشتراك الخاصة بك.")
                    return redirect('subscribe')

                # Check if the user is trying to subscribe to the same plan again
                if user_profile.subscription_plan == new_subscription_plan:
                    messages.error(request, "أنت بالفعل مشترك في هذه الخطة.")
                    return redirect('subscribe')

            # Calculate subscription duration
            duration = calculate_subscription_duration(new_subscription_plan)
            if duration is None:
                messages.error(request, f"المدة غير معروفة لخطة الاشتراك '{subscription_plan_name}'.")
                return redirect('subscribe')
            
            # Calculate subscription end date
            start_date = timezone.now()
            end_date = start_date + duration
            
            # Save subscription details to user's profile
            user_profile.subscription_start_date = start_date
            user_profile.subscription_end_date = end_date
            user_profile.subscription_plan = new_subscription_plan
            user_profile.save()
            
            # Send email confirmation
            subject = 'تأكيد الاشتراك'
            html_message = render_to_string('users/subscription_confirmation.html', {'user': request.user})
            plain_message = strip_tags(html_message)  # Strip HTML tags for plain text email
            from_email = 'adaladjango@gmail.com'  # Replace with your email address
            to_email = request.user.email
            send_mail(subject, plain_message, from_email, [to_email], html_message=html_message)
            
             # Redirect users to their respective home pages based on their role
            if request.user.role == User.Role.LAWYER:
                messages.success(request, f"تم الاشتراك بنجاح في الخطة '{subscription_plan_name}'.")
                return redirect('lawyerhome')
                
            elif request.user.role == User.Role.NOTARY:
                messages.success(request, f"تم الاشتراك بنجاح في الخطة '{subscription_plan_name}'.")
                return redirect('notaryhome')
            else:
                return redirect('home2')
        else:
            messages.error(request, "الرجاء اختيار خطة الاشتراك.")
            return redirect('subscribe')
    else:
        # Retrieve all subscription plans from the database
        subscription_plans = SubscriptionPlan.objects.all()
        
       # Calculate remaining time for the user's subscription
        remaining_time = None
        if request.user.is_authenticated and request.user.profile.subscription_end_date:
            current_time = timezone.now()
            subscription_end_date = request.user.profile.subscription_end_date
            if current_time < subscription_end_date:
                remaining_time = subscription_end_date - current_time
                
                # Convert remaining time to a user-friendly format
                days = remaining_time.days
                hours, remainder = divmod(remaining_time.seconds, 3600)
                minutes, seconds = divmod(remainder, 60)
                remaining_time = f"{days} days, {hours} hours, {minutes} minutes, {seconds} seconds"
        
        return render(request, 'users/subscribe.html', {'subscription_plans': subscription_plans, 'remaining_time': remaining_time})
def calculate_subscription_duration(subscription_plan):
    if subscription_plan.duration == 'month':
        return timedelta(days=30)
    elif subscription_plan.duration == '1_year':
        return timedelta(days=365)  # Assuming 1 year is 365 days
    elif subscription_plan.duration == '6_months':
        return timedelta(days=30*6)  # 6 months duration
    else:
        return None  # Unknown duration



################################################## My simple message system #################################################
@login_required
def inbox(request):
    # Get distinct sender IDs who have sent messages to the current user
    senders_ids = Message.objects.filter(recipient=request.user).values_list('sender_id', flat=True).distinct()
    # Get User objects corresponding to the sender IDs
    senders = User.objects.filter(id__in=senders_ids)
    return render(request, 'users/inbox.html', {'senders': senders})

@login_required
def ninbox(request):
    # Get distinct sender IDs who have sent messages to the current user
    senders_ids = Message.objects.filter(recipient=request.user).values_list('sender_id', flat=True).distinct()
    # Get User objects corresponding to the sender IDs
    senders = User.objects.filter(id__in=senders_ids)
    return render(request, 'users/ninbox.html', {'senders': senders})




@login_required
def conversation(request, sender_id):
    sender = User.objects.get(pk=sender_id)
    messages = Message.objects.filter(sender_id=sender_id, recipient=request.user) | Message.objects.filter(sender_id=request.user, recipient_id=sender_id)
    messages = messages.order_by('timestamp')  # Order messages by timestamp
    return render(request, 'users/conversation.html', {'messages': messages, 'sender': sender})


@login_required
def send_message(request, sender_id):
    if request.method == 'POST':
        recipient = User.objects.get(pk=sender_id)
        content = request.POST.get('content')
        file = request.FILES.get('file')  # Get uploaded file
        
        # Create a new Message instance using the sender_id, recipient, content, and file
        message = Message.objects.create(sender_id=request.user, recipient=recipient, content=content)
        if file:
            message.file.save(file.name, ContentFile(file.read()))  # Save the file
        
        return redirect('conversation', sender_id=sender_id)
    return redirect('inbox')





def delete_message(request, message_id):
    if request.method == 'POST':
        message = get_object_or_404(Message, pk=message_id)
        sender_id = None  # Initialize sender_id variable
        
        # Check if the user is the sender or recipient of the message
        if request.user == message.sender_id or request.user == message.recipient:
            if request.user == message.sender_id:
                sender_id = message.recipient_id
            else:
                sender_id = message.sender_id
            
            message.delete()
            messages.success(request, 'تم حذف الرسالة بنجاح!')
        else:
            messages.error(request, 'ليس لديك الإذن لحذف هذه الرسالة.')
    
    # Redirect back to the conversation page with the appropriate sender_id
    if sender_id is not None:
        return redirect(reverse('conversation', kwargs={'sender_id': sender_id}))
    else:
        return redirect('inbox')  # Redirect to inbox if sender_id is not found
    




# CRUD Functionality
# Create


# ADMIN

def news_detail(request, pk):
    news = get_object_or_404(News, pk=pk)
    return render(request, 'users/detail.html', {'news': news})


@staff_member_required
def create_news(request):
    if request.method == 'POST':
        form = NewsForm(request.POST, request.FILES)  # Pass request.FILES to handle uploaded files
        if form.is_valid():
            news = form.save(commit=False)
            news.author = request.user
            news.save()
            return redirect('news_detail', pk=news.pk)
    else:
        form = NewsForm()
    return render(request, 'users/create.html', {'form': form})
@staff_member_required
def update_news(request, pk):
    news = get_object_or_404(News, pk=pk)
    if request.method == 'POST':
        form = NewsForm(request.POST, request.FILES, instance=news)
        if form.is_valid():
            news = form.save(commit=False)
            news.author = request.user
            news.save()
            return redirect('news_detail', pk=news.pk)
    else:
        form = NewsForm(instance=news)
    return render(request, 'users/update.html', {'form': form})

# Delete
@staff_member_required
def delete_news(request, pk):
    news = get_object_or_404(News, pk=pk)
    news.delete()
    return redirect('news_list')

@staff_member_required
def news_list(request):
    news_articles = News.objects.all()
    return render(request, 'users/news_list.html', {'news_articles': news_articles})

@staff_member_required
def admin_dashboard(request):
    return render(request, 'users/dashboard.html')



def registerall(request):
    # You can pass any necessary context data here

    return render(request, 'users/registerall.html')
def homesearch_results(request):
    if request.method == "GET":
        assistant_type = request.GET.get("assistant_type")
        practice_area = request.GET.get("practice_area")

        if assistant_type == "LAWYER":
            assistants = Lawyer.objects.filter(practice_areas__icontains=practice_area)
        elif assistant_type == "NOTARY":
            assistants = Notary.objects.filter(practice_areas__icontains=practice_area)
        else:
            assistants = []

        # Define the context dictionary
        context = {
            'assistants': assistants,
            # You can add more key-value pairs to the context as needed
        }

        return render(request, 'users/homesearch_results.html', context)
    else:
        return render(request, 'users/homesearch_results.html', {})
from docx import Document
from io import BytesIO
from datetime import datetime
from docx.shared import Pt

def generate_word_document(request):
    if request.method == 'POST':
        # Initialize a new Document object
        doc = Document()

        # Get data from the form and store it in a dictionary
        data = {
            'vendeur': request.POST.get('vendeur', ''), # البائع
            'vdate': request.POST.get('vdate', ''), # تاريخ ازدياد البائع
            'vlieu': request.POST.get('vlieu', ''), # مكان ازدياد البائع
            'vnumcarte': request.POST.get('vnumcarte', ''), # رقم بطاقة تعريف  البائع
            'vdatecarte': request.POST.get('vdatecarte', ''), # تاريخ اصدار بطاقة البائع
            'vlieucarte': request.POST.get('vlieucarte', ''),   #مكان اصدار بطاقة البائع
            'vadresse': request.POST.get('vadresse', ''),# عنوان البائع 
            'acheteur': request.POST.get('acheteur', ''),#المشتري
            'adate': request.POST.get('adate', ''),# تاريخ ازدياد المشتري
            'alieu': request.POST.get('alieu', ''),# مكان ازدياد المشتري
            'anumcarte': request.POST.get('anumcarte', ''),# رقم بطاقة تعريف المشتري
            'adatecarte': request.POST.get('adatecarte', ''),
            'alieucarte': request.POST.get('alieucarte', ''),
            'aadresse': request.POST.get('aadresse', ''),
            'temoin1': request.POST.get('temoin1', ''),
            'tdre1': request.POST.get('tdre1', ''),
            'tlieu1': request.POST.get('tlieu1', ''),
            'tnumcarte1': request.POST.get('tnumcarte1', ''),
            'tdatecarte1': request.POST.get('tdatecarte1', ''),
            'tlieucarte1': request.POST.get('tlieucarte1', ''),
            'tadresse1': request.POST.get('tadresse1', ''),
            'temoin2': request.POST.get('temoin2', ''),
            'tdre2': request.POST.get('tdre2', ''),
            'tlieu2': request.POST.get('tlieu2', ''),
            'tnumcarte2': request.POST.get('tnumcarte2', ''),
            'tdatecarte2': request.POST.get('tdatecarte2', ''),
            'tlieucarte2': request.POST.get('tlieucarte2', ''),
            'tadresse2': request.POST.get('tadresse2', ''),
            'logement': request.POST.get('logement', '')
        }

        # Accessing styles via the Document instance
        style_title = doc.styles['Heading 1']
        style_title.font.size = Pt(16)
        style_title.font.bold = True

        style_paragraph = doc.styles['Normal']
        style_paragraph.font.size = Pt(12)
        
        current_date = datetime.now().strftime("%Y-%m-%d")
      
        # Add the main text to the document
        text = (         
            f" عقد بيع عقار \n"
            f" {data['vendeur']} - {data['acheteur']}    \n"
            f" بتاريخ:{current_date} \n"
            "محرر عرفي يتضمن عقد بيع\n"
            f"أمام الكاتب العمومي {data['vendeur']} الموقع أدناه\n"
            "حضر لدينا\n"
            f"الطرف البائع السيد(ة): {data['vendeur']} المولود بتاريخ:{data['vdate']} بـ:{data['vlieu']}  الحامل لبطاقة تعريف رقم:{data['vnumcarte']} الصادرة بتاريخ: {data['vdatecarte']} عن: {data['vlieucarte']} الساكن: {data['vadresse']} جزائري الجنسية\n"
            "الـذي (التي – ان-ن) صــرح (ت-وا) حــال صحتــه (ا-م) وجـواز أمـره (ا-م) شرعا وقانونا مع تمتعه (ا-م) بالأهلية المدنية والقانونيــة للتصــرف بمـوجـب هـــذا التصريـــح , بأنـــه (ا-م) قـــد باع (ا- ت- وا) إلـى\n"
            f"الطرف الشاري السيد(ة): {data['acheteur']} المولود بتاريخ:{data['adate']} بـ:{data['alieu']} الحامل لبطاقة التعريف رقم:{data['anumcarte']} الصادرة بتاريخ: {data['adatecarte']} عن: {data['alieucarte']} الساكن:  {data['aadresse']} جزائري الجنسية\n"
            f"العقار الذي بيانه //\n"
            f"تعيين العقار:\nوقد آل إلى البائع {data['logement']} عن طريق: \n"
            "الثمن: حدد ثمن البيع بمبلغ متفق عليه مقبوضة نقدا، بيعا صحيحا لا رجعة فيه تم بمحضر وشهادة الشاهدين\n"
            "الالتزام: التزام البائع بنقل حق الانتفاع للعقار المبين أعلاه إلى الطرف الشاري ليصبح يتمتع بالحق والحرية الكاملة للتصرف بالعقار كما يشاء، كما صرح بأن العقار المباع ليس قيد رهن أو على الشياع مع الغير وأكد أنه خال من كل التبعيات\n"
            "الموطن: أجل تنفيذ هذا التصريح وتوابعه أختار كل طرف التصريح موطنه القانوني بمحل سكناه المذكورة أعلاه.\n"
        )
        text += ( f"\n"
            f"الطرف الشاهد1 السيد(ة): {data.get('temoin1', '')} المولود بتاريخ:{data.get('tdre1', '')} بـ:{data.get('tlieu1', '')} الحامل لبطاقة التعريف رقم:{data.get('tnumcarte1', '')} الصادرة بتاريخ: {data.get('tdatecarte1', '')} عن: {data.get('tlieucarte1', '')} الساكن: {data.get('tadresse1', '')} جزائري الجنسية\n"
            f"الطرف الشاهد2 السيد(ة): {data.get('temoin2', '')} المولود بتاريخ:{data.get('tdre2', '')} بـ:{data.get('tlieu2', '')} الحامل لبطاقة التعريف رقم:{data.get('tnumcarte2', '')} الصادرة بتاريخ: {data.get('tdatecarte2', '')} عن: {data.get('tlieucarte2', '')} الساكن: {data.get('tadresse2', '')} جزائري الجنسية\n"
            f"اثباتا لما ذكر: حـرر و تـم هــذا التصريـح ووقـع بمكتب {data.get('vendeur', '')}بتاريخ  كاتب عمومي   : {current_date}\n"
            "وبعـد التــلاوة التـي تمـت امـضي التصريـح مـن طـرف المصرحيـن والشهـود في اليوم والسنة المذكورين أعلاه.\n"
            "يتكون هذا التصريح من....صفحة مخطوطا على....فراغات بدون تغيير أو زيـادة في الحروف أو تشطيب أمضي التصريح\n"
            "يليه الإمضاءات\n"
            "صفة المحرر                      البائع                      الشاري                                  الشاهد الاول                             الشاهد الثاني"
        )
        paragraph = doc.add_paragraph(text)
        paragraph.alignment = 2 

        # Save the document to a BytesIO object
        output = BytesIO()
        doc.save(output)
        output.seek(0)

        # Prepare response as a downloadable Word document
        response = HttpResponse(
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = 'attachment; filename=contract_document.docx'
        response.write(output.getvalue())
        return response

    return render(request, 'users/vente.html')
def vente_cycle(request):
    if request.method == 'POST':
        # Initialize a new Document object
        doc = Document()

        # Get data from the form and store it in a dictionary
        data = {
            'vendeur': request.POST.get('vendeur', ''),
            'vdate': request.POST.get('vdate', ''),
            'vlieu': request.POST.get('vlieu', ''),
            'vnumcarte': request.POST.get('vnumcarte', ''),
            'vdatecarte': request.POST.get('vdatecarte', ''),
            'vlieucarte': request.POST.get('vlieucarte', ''),
            'vadresse': request.POST.get('vadresse', ''),
            'acheteur': request.POST.get('acheteur', ''),
            'adate': request.POST.get('adate', ''),
            'alieu': request.POST.get('alieu', ''),
            'anumcarte': request.POST.get('anumcarte', ''),
            'adatecarte': request.POST.get('adatecarte', ''),
            'alieucarte': request.POST.get('alieucarte', ''),
            'aadresse': request.POST.get('aadresse', ''),
            'source': request.POST.get('source', ''),
            'typec': request.POST.get('typec', ''),
            'colorc': request.POST.get('colorc', ''),
            'numc': request.POST.get('numc', ''),
        }

        # Accessing styles via the Document instance
        style_title = doc.styles['Heading 1']
        style_title.font.size = Pt(16)
        style_title.font.bold = True

        style_paragraph = doc.styles['Normal']
        style_paragraph.font.size = Pt(12)
        
        current_date = datetime.now().strftime("%Y-%m-%d")
      
        # Add the main text to the document
        text = (         
            "رقم: 3_2024\n"
            f" بتاريخ:{current_date} \n"       
            "تصريح بيع دراجة نارية\n"
            f"البائع السيد(ة): {data['vendeur']} المولود بتاريخ: {data['vdate']} بـ: {data['vlieu']} الحامل لبطاقة التعريف  رقم: {data['vnumcarte']} الصادرة بتاريخ: {data['vdatecarte']} عن: {data['vlieucarte']} الساكن: {data['vadresse']}، جزائري الجنسية\n"
            "أصرح بشرفي أني بعت دراجة نارية ذات المواصفات التالية: \n"
            f"النوع: {data['typec']}       اللون: {data['colorc']}       رقم الهيكل: {data['numc']}\n"
            f"الشاري السيد(ة): {data['acheteur']} المولود بتاريخ: {data['adate']} بـ: {data['alieu']} الحامل لبطاقة التعريف رقم: {data['anumcarte']} الصادرة بتاريخ: {data['adatecarte']} عن: {data['alieucarte']} الساكن: {data['aadresse']}، جزائري الجنسية\n"
            "وقد آلت إلى البائع عن طريق: [اكتب طريقة الاتصال] في: 11/05/2024\n"
            "صفة المحرر                                                 البائع                                                 الشاري\n"
        )

        paragraph = doc.add_paragraph(text)
        paragraph.alignment = 2 

        # Save the document to a BytesIO object
        output = BytesIO()
        doc.save(output)
        output.seek(0)

        # Prepare response as a downloadable Word document
        response = HttpResponse(
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = 'attachment; filename=contract_document.docx'
        response.write(output.getvalue())
        return response

    return render(request, 'users/ventecycle.html')
def vente_car(request):
    if request.method == 'POST':
        # Initialize a new Document object
        doc = Document()

        # Get data from the form and store it in a dictionary
        data = {
            'vendeur': request.POST.get('vendeur', ''),
            'vdate': request.POST.get('vdate', ''),
            'vlieu': request.POST.get('vlieu', ''),
            'vnumcarte': request.POST.get('vnumcarte', ''),
            'vdatecarte': request.POST.get('vdatecarte', ''),
            'vlieucarte': request.POST.get('vlieucarte', ''),
            'vadresse': request.POST.get('vadresse', ''),
            'acheteur': request.POST.get('acheteur', ''),
            'adate': request.POST.get('adate', ''),
            'alieu': request.POST.get('alieu', ''),
            'anumcarte': request.POST.get('anumcarte', ''),
            'adatecarte': request.POST.get('adatecarte', ''),
            'alieucarte': request.POST.get('alieucarte', ''),
            'aadresse': request.POST.get('aadresse', ''),
            'source': request.POST.get('source', ''),
            'typec': request.POST.get('typec', ''),
            'colorc': request.POST.get('colorc', ''),
            'numc': request.POST.get('numc', ''),
        }

        # Accessing styles via the Document instance
        style_title = doc.styles['Heading 1']
        style_title.font.size = Pt(16)
        style_title.font.bold = True

        style_paragraph = doc.styles['Normal']
        style_paragraph.font.size = Pt(12)
        
        current_date = datetime.now().strftime("%Y-%m-%d")
      
        # Add the main text to the document
        text = (         
            "رقم: 3_2024\n"
            f" بتاريخ:{current_date} \n"                 "تصريح بيع مركبة \n"
            f"البائع السيد(ة): {data['vendeur']} المولود بتاريخ: {data['vdate']} بـ: {data['vlieu']} الحامل لبطاقة التعريف  رقم: {data['vnumcarte']} الصادرة بتاريخ: {data['vdatecarte']} عن: {data['vlieucarte']} الساكن: {data['vadresse']}، جزائري الجنسية\n"
            "أصرح بشرفي أني بعت  مركبة ذات المواصفات التالية: \n"
            f"النوع: {data['typec']}       اللون: {data['colorc']}       رقم الهيكل: {data['numc']}\n"
            f"الشاري السيد(ة): {data['acheteur']} المولود بتاريخ: {data['adate']} بـ: {data['alieu']} الحامل لبطاقة التعريف رقم: {data['anumcarte']} الصادرة بتاريخ: {data['adatecarte']} عن: {data['alieucarte']} الساكن: {data['aadresse']}، جزائري الجنسية\n"
            f"وقد آلت إلى البائع عن طريق: {data['source']} في:{current_date} //\n"
            "صفة المحرر                                                 البائع                                                 الشاري\n"
        )

        paragraph = doc.add_paragraph(text)
        paragraph.alignment = 2 
        # Save the document to a BytesIO object
        output = BytesIO()
        doc.save(output)
        output.seek(0)
        print(doc.styles)
        print(style_title)
        print(style_paragraph)

        # Prepare response as a downloadable Word document
        response = HttpResponse(
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = 'attachment; filename=contract_document.docx'
        response.write(output.getvalue())
        return response

    return render(request, 'users/ventecar.html')
def dette(request):
    if request.method == 'POST':
        # Initialize a new Document object
        doc = Document()

        # Get data from the form and store it in a dictionary
        data = {
            'vendeur': request.POST.get('vendeur', ''), # البائع
            'vdate': request.POST.get('vdate', ''), # تاريخ ازدياد البائع
            'vlieu': request.POST.get('vlieu', ''), # مكان ازدياد البائع
            'vnumcarte': request.POST.get('vnumcarte', ''), # رقم بطاقة تعريف  البائع
            'vdatecarte': request.POST.get('vdatecarte', ''), # تاريخ اصدار بطاقة البائع
            'vlieucarte': request.POST.get('vlieucarte', ''),   #مكان اصدار بطاقة البائع
            'vadresse': request.POST.get('vadresse', ''),# عنوان البائع 
            'acheteur': request.POST.get('acheteur', ''),#المشتري
            'adate': request.POST.get('adate', ''),# تاريخ ازدياد المشتري
            'alieu': request.POST.get('alieu', ''),# مكان ازدياد المشتري
            'anumcarte': request.POST.get('anumcarte', ''),# رقم بطاقة تعريف المشتري
            'adatecarte': request.POST.get('adatecarte', ''),
            'alieucarte': request.POST.get('alieucarte', ''),
            'aadresse': request.POST.get('aadresse', ''),
            'temoin1': request.POST.get('temoin1', ''),
            'tdre1': request.POST.get('tdre1', ''),
            'tlieu1': request.POST.get('tlieu1', ''),
            'tnumcarte1': request.POST.get('tnumcarte1', ''),
            'tdatecarte1': request.POST.get('tdatecarte1', ''),
            'tlieucarte1': request.POST.get('tlieucarte1', ''),
            'tadresse1': request.POST.get('tadresse1', ''),
            'temoin2': request.POST.get('temoin2', ''),
            'tdre2': request.POST.get('tdre2', ''),
            'tlieu2': request.POST.get('tlieu2', ''),
            'tnumcarte2': request.POST.get('tnumcarte2', ''),
            'tdatecarte2': request.POST.get('tdatecarte2', ''),
            'tlieucarte2': request.POST.get('tlieucarte2', ''),
            'tadresse2': request.POST.get('tadresse2', ''),
            'logement': request.POST.get('logement', '')
        }

        # Accessing styles via the Document instance
        style_title = doc.styles['Heading 1']
        style_title.font.size = Pt(16)
        style_title.font.bold = True

        style_paragraph = doc.styles['Normal']
        style_paragraph.font.size = Pt(12)
        
        current_date = datetime.now().strftime("%Y-%m-%d")
      
        # Add the main text to the document
        text = ( 
    f"اعتراف بدين\n"
    f"{data['vendeur']} - {data['acheteur']}\n"
    f"بتاريخ: {current_date}\n"
    "محرر عرفي يتضمن عقد بيع\n"
    f"أمام الكاتب العمومي {data['vendeur']} الموقع أدناه\n"
    "حضر لدينا\n"
    f"المدين: {data['vendeur']} المولود بتاريخ: {data['vdate']} بـ: {data['vlieu']} الحامل لبطاقة تعريف رقم: {data['vnumcarte']} الصادرة بتاريخ: {data['vdatecarte']} عن: {data['vlieucarte']} الساكن: {data['vadresse']} جزائري الجنسية\n"
    f"{data['logement']} الـذي (التي) شهد على نفسه أنه مدين بمبلغ مالي يقدر ب\n"
    "تسلمه خارج مجلس العقد وبعيداً عن رؤية الموثق الموقع أسفله من:\n"
    f"الدائن السيد(ة): {data['acheteur']} المولود بتاريخ: {data['adate']} بـ: {data['alieu']} الحامل لبطاقة التعريف رقم: {data['anumcarte']} الصادرة بتاريخ: {data['adatecarte']} عن: {data['alieucarte']} الساكن: {data['aadresse']} جزائري الجنسية\n"
    ": الشروط والالتزامات\n"
    "   في حالة وفاة المدين فإن الدين ينتقل إلى ورثته من بعده وبالتالي فهو ملزم برد الدين إلى الدائن المذكور أعلاه في الأجل المحد\n"
    "   كذلك يستوي الأمر في حالة وفاة الدائن فإن ورثته يحلون محله في مطالبة المدين بمبلغ الدين المصرح به أمام الموثق في الأجال المحددة له\n"
    "   يتعين على المدين أن يرد مبلغ الدين للدائن أعلاه أو لورثته وقد نبه الموثق الطرفان عند الوفاء بضرورة إبرام عقد إبراء الدين في محرر رسمي أمام الموثق\n"
    "   في حالة عدم الوفاء من جهة المدين فإن للدائن الحق في اللجوء إلى تطبيق النسخة التنفيذية التي يحوزها قانونًا والتي تنفذ عند الضرورة جبرًا على أموال وممتلكات مدينه\n"
    "  ومن جهته يلتزم الموثق الموقع أسفله بتسليم نسخة عادية واحدة إلى المدين ونسخة تنفيذية واحدة إلى الدائن يستعمل حقه بها متى كان هناك اختلال بعدم الوفاء من قبل المدين //\n"
    "الثمن: حدد ثمن البيع بمبلغ متفق عليه مقبوضة نقدًا، بيعًا صحيحًا لا رجعة فيه تم بمحضر وشهادة الشاهدين\n"
    "تم هذا العقد بحضور الشاهدين وهما:\n"
    f"الطرف الشاهد1 السيد(ة): {data.get('temoin1', '')} المولود بتاريخ: {data.get('tdre1', '')} بـ: {data.get('tlieu1', '')} الحامل لبطاقة التعريف رقم: {data.get('tnumcarte1', '')} الصادرة بتاريخ: {data.get('tdatecarte1', '')} عن: {data.get('tlieucarte1', '')} الساكن: {data.get('tadresse1', '')} جزائري الجنسية\n"
    f"الطرف الشاهد2 السيد(ة): {data.get('temoin2', '')} المولود بتاريخ: {data.get('tdre2', '')} بـ: {data.get('tlieu2', '')} الحامل لبطاقة التعريف رقم: {data.get('tnumcarte2', '')} الصادرة بتاريخ: {data.get('tdatecarte2', '')} عن: {data.get('tlieucarte2', '')} الساكن: {data.get('tadresse2', '')} جزائري الجنسية\n"
    f"اثباتًا لما ذكر: حرر وتم هذا التصريح ووقع بمكتب {data.get('vendeur', '')} بتاريخ كاتب عمومي: {current_date}\n"
    "وبعد التلاوة التي تمت امضي التصريح من طرف المصرحين والشهود في اليوم والسنة المذكورين أعلاه.\n"
    "يتكون هذا التصريح من....صفحة مخطوطة على....فراغات بدون تغيير أو زيادة في الحروف أو تشطيب أمضي التصريح\n"
    "يليه الإمضاءات\n"
    "صفة المحرر الدائن المدين الشاهد الأول الشاهد الثاني"
)



        paragraph = doc.add_paragraph(text)
        paragraph.alignment = 2 

        # Save the document to a BytesIO object
        output = BytesIO()
        doc.save(output)
        output.seek(0)

        # Prepare response as a downloadable Word document
        response = HttpResponse(
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = 'attachment; filename=contract_document.docx'
        response.write(output.getvalue())
        return response

    return render(request, 'users/dette.html')

@login_required
def npaiement(request):
    if request.method == "POST":
        try:
            # Get form data
            nomclient1 = request.POST.get('nomclient1')
            nomclient2 = request.POST.get('nomclient2')
            datep = request.POST.get('datep')
            budjetp = request.POST.get('budjetp')
            details = request.POST.get('details')
            type = request.POST.get('type')
            adressep = request.POST.get('adressep')
           
            # Get the authenticated user's notary instance
            notary = request.user.notary

            # Create a new npaiements instance and associate it with the notary
            paiement = npaiements.objects.create (
                notary=notary,
                fpartie=nomclient1,
                spartie=nomclient2,
                date=datep,
                prix=budjetp,
                remarques=details,
                type=type,
                address=adressep
            )
            messages.success(request, 'Judgment added successfully!')
            return redirect('tablenpaiement')
        except Exception as e:
            # If an error occurs during database operation, display an error message
            messages.error(request, f'Error: {e}')
            return redirect('npaiements')  # Redirect to the same page
    else:
        # Handle GET request to render the form
        notary = request.user.notary
        clients = notaryclients.objects.filter(notary=notary)
        
        # Initialize nomclient with the first client's name if available
        initial_nomclient = f"{clients.first().Fname} {clients.first().Sname}" if clients.exists() else None

        return render(request, "users/npaiements.html", {'initial_nomclient': initial_nomclient, 'clients': clients})
@login_required  # Assure que l'utilisateur est connecté
def tablenpaiement(request):
    # Récupérer l'utilisateur connecté
    notary = request.user

    # Filtrer les affaires liées à l'avocat connecté
    npaiement = npaiements.objects.filter(notary=notary)

    context = {'npaiements': npaiement}
    return render(request, "users/tablenpaiement.html", context)
def delete_npaiement(request, cid):
    try:
        paiement_instance = npaiements.objects.get(id=cid)
        paiement_instance.delete()
        messages.success(request, '! تمت الحذف بنجاح')
    except Seance.DoesNotExist:
        messages.error(request, 'العنصر المطلوب غير موجود')
    except Exception as e:
        messages.error(request, f'خطأ: {e}')
    
    # Redirect back to the same page
    return HttpResponseRedirect(request.META.get('HTTP_REFERER', '/'))

@login_required
def update_npaiement(request, npaiement_id):
    npaiement = get_object_or_404(npaiements, id=npaiement_id)

    if request.method == "POST":
        # Define the form directly here
        class NpaiementsForm(forms.ModelForm):
            class Meta:
                model = npaiements
                fields = ['notary', 'fpartie', 'spartie', 'date', 'type', 'address', 'prix', 'remarques']
        
        form = NpaiementsForm(request.POST, instance=npaiement)
        if form.is_valid():
            form.save()
            messages.success(request, "تم تحديث البيانات بنجاح.")
            return redirect('tablenpaiement')  # Redirect to the list view of npaiements after update
    else:
        # Define the form directly here
        class NpaiementsForm(forms.ModelForm):
            class Meta:
                model = npaiements
                fields = ['notary', 'fpartie', 'spartie', 'date', 'type', 'address', 'prix', 'remarques']

        form = NpaiementsForm(instance=npaiement)

    return render(request, 'users/update_npaiement.html', {'form': form, 'npaiement': npaiement})
@login_required
def manage_clients_notary(request):
    if request.user.role != User.Role.NOTARY:
        messages.error(request, "ليس لديك الصلاحية للوصول إلى هذه الصفحة.")
        return redirect('home2')

    notary = request.user.notary
    if notary is None:
        messages.error(request, "You are not registered as a notary.")
        return redirect('notaryhome')

    if request.method == 'POST':
        selected_clients = request.POST.getlist('selected_clients')
        subscription = request.user.profile.subscription_plan  # Access subscription plan through the user's profile

        if subscription is None:
            messages.error(request, "ليس لديك خطة اشتراك نشطة.")
            return redirect('notaryhome')

        max_accepted_clients = subscription.max_accepted_clients
        accepted_count = notary.accepted_clients.count()

        non_existent_clients = []

        for client_id in selected_clients:
            try:
                client = User1.objects.get(pk=client_id)
            except User1.DoesNotExist:
                non_existent_clients.append(client_id)
                continue

            if NotaryClient.objects.filter(notary=notary, user1=client).exists():
                messages.warning(request, f"العميل '{client.username}' مقبول بالفعل.")
            elif max_accepted_clients is not None and accepted_count >= max_accepted_clients:
                messages.error(request, "لقد وصلت إلى الحد الأقصى لعدد العملاء المقبولين.")
            else:
                NotaryClient.objects.create(notary=notary, user1=client)
                messages.success(request, f"تم قبول العميل '{client.username}' بنجاح.")

                # Send a message to the client
                content = "تم قبولك من قبل الموثق."
                Message.objects.create(sender_id=request.user, recipient=client, content=content)
                accepted_count += 1

        if non_existent_clients:
            messages.error(request, f"The following clients do not exist: {', '.join(non_existent_clients)}")

        return redirect('manage_clients_notary')
    else:
        message_senders = Message.objects.filter(recipient=request.user)
        potential_clients = User1.objects.filter(id__in=message_senders.values('sender_id'))
        context = {'potential_clients': potential_clients}
        return render(request, 'users/manage_clients_notary.html', context)